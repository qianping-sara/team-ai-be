import os
from docx import Document as DocxDocument
from PyPDF2 import PdfReader
from datetime import datetime
import uuid
from typing import List, Tuple, Dict, Any
from models.document_models import Document, DocumentSection, DocumentContent, ContentType, DocumentStatus
from database.mongo_client import documents, document_sections, document_contents
import re
from io import BytesIO

class DocumentProcessor:
    def __init__(self, filename: str, file_content: bytes):
        self.filename = filename
        self.file_content = file_content
        self.file_type = self._get_file_type()
        self.document_id = str(uuid.uuid4())
        
    def _get_file_type(self) -> str:
        """根据文件扩展名判断文件类型"""
        ext = os.path.splitext(self.filename)[1].lower()
        type_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.txt': 'text/plain'
        }
        return type_map.get(ext, 'application/octet-stream')

    def process_and_save(self) -> str:
        """处理文档并保存到MongoDB"""
        try:
            # 创建文档记录
            doc = Document(self.filename, self.file_type)
            doc.data['_id'] = self.document_id
            doc.data['status'] = DocumentStatus.PROCESSING.value
            documents.insert_one(doc.data)

            try:
                # 处理文档内容
                if self.file_type == 'application/pdf':
                    self._process_pdf()
                elif self.file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
                    self._process_docx()
                elif self.file_type == 'application/msword':
                    # 对于 .doc 文件，创建一个默认章节并添加提示信息
                    section = self._create_default_section()
                    content = DocumentContent(
                        self.document_id,
                        section['_id'],
                        ContentType.TEXT
                    )
                    content.set_text_content("注意：暂不支持直接处理 .doc 格式文件，请将文件转换为 .docx 格式后重新上传。")
                    content.data['_id'] = str(uuid.uuid4())
                    content.data['order'] = 10
                    document_contents.insert_one(content.data)
                else:
                    self._process_text()

                # 更新文档状态
                documents.update_one(
                    {'_id': self.document_id},
                    {'$set': {
                        'status': DocumentStatus.PROCESSED.value,
                        'last_modified': datetime.utcnow()
                    }}
                )

            except Exception as e:
                # 更新错误状态
                documents.update_one(
                    {'_id': self.document_id},
                    {'$set': {
                        'status': DocumentStatus.ERROR.value,
                        'error_message': str(e)
                    }}
                )
                raise

            return self.document_id

        except Exception as e:
            raise Exception(f"处理文档时出错: {str(e)}")

    def _process_docx(self):
        """处理Word文档"""
        try:
            doc = DocxDocument(BytesIO(self.file_content))
            current_section = None
            section_stack = []
            last_level = 0

            for paragraph in doc.paragraphs:
                # 检查是否是章节标题
                if paragraph.style.name.startswith('Heading'):
                    level = int(paragraph.style.name.replace('Heading ', ''))
                    
                    # 处理章节层级
                    while section_stack and last_level >= level:
                        section_stack.pop()
                        last_level -= 1

                    # 创建新章节
                    parent_id = section_stack[-1]['_id'] if section_stack else None
                    section = DocumentSection(
                        self.document_id,
                        paragraph.text,
                        level,
                        parent_id
                    )
                    section.data['_id'] = str(uuid.uuid4())
                    section.data['order'] = self._get_next_section_order(parent_id)
                    section.data['section_number'] = self._generate_section_number(section_stack, section.data['order'])
                    
                    # 保存章节
                    document_sections.insert_one(section.data)
                    current_section = section.data
                    section_stack.append(current_section)
                    last_level = level
                
                # 处理段落内容
                elif paragraph.text.strip():
                    if not current_section:
                        # 如果没有章节，创建默认章节
                        current_section = self._create_default_section()
                    
                    # 创建文本内容
                    content = DocumentContent(
                        self.document_id,
                        current_section['_id'],
                        ContentType.TEXT
                    )
                    content.set_text_content(paragraph.text)
                    content.data['_id'] = str(uuid.uuid4())
                    content.data['order'] = self._get_next_content_order(current_section['_id'])
                    
                    # 保存内容
                    document_contents.insert_one(content.data)

            # 处理表格
            for table in doc.tables:
                if not current_section:
                    current_section = self._create_default_section()

                headers = []
                rows = []
                
                # 提取表格数据
                for i, row in enumerate(table.rows):
                    row_data = [cell.text for cell in row.cells]
                    if i == 0:
                        headers = row_data
                    else:
                        rows.append(row_data)

                # 创建表格内容
                content = DocumentContent(
                    self.document_id,
                    current_section['_id'],
                    ContentType.TABLE
                )
                content.set_table_content(headers, rows)
                content.data['_id'] = str(uuid.uuid4())
                content.data['order'] = self._get_next_content_order(current_section['_id'])
                
                # 保存内容
                document_contents.insert_one(content.data)

        except Exception as e:
            raise Exception(f"处理Word文档时出错: {str(e)}")

    def _create_default_section(self) -> Dict[str, Any]:
        """创建默认章节"""
        section = DocumentSection(
            self.document_id,
            "未分类内容",
            0
        )
        section.data['_id'] = str(uuid.uuid4())
        section.data['order'] = self._get_next_section_order(None)
        section.data['section_number'] = str(section.data['order'])
        document_sections.insert_one(section.data)
        return section.data

    def _get_next_section_order(self, parent_id: str = None) -> int:
        """获取下一个章节序号"""
        last_section = document_sections.find_one(
            {'document_id': self.document_id, 'parent_id': parent_id},
            sort=[('order', -1)]
        )
        return (last_section['order'] + 10) if last_section else 10

    def _get_next_content_order(self, section_id: str) -> int:
        """获取下一个内容序号"""
        last_content = document_contents.find_one(
            {'section_id': section_id},
            sort=[('order', -1)]
        )
        return (last_content['order'] + 10) if last_content else 10

    def _generate_section_number(self, section_stack: List[Dict[str, Any]], order: int) -> str:
        """生成章节编号"""
        if not section_stack:
            return str(order // 10)
        
        parent_number = section_stack[-1]['section_number']
        return f"{parent_number}.{order // 10}"

    def _process_pdf(self):
        """处理PDF文档"""
        try:
            reader = PdfReader(BytesIO(self.file_content))
            section = self._create_default_section()
            
            for i, page in enumerate(reader.pages):
                content = DocumentContent(
                    self.document_id,
                    section['_id'],
                    ContentType.TEXT
                )
                content.set_text_content(page.extract_text())
                content.data['_id'] = str(uuid.uuid4())
                content.data['order'] = (i + 1) * 10
                document_contents.insert_one(content.data)
        except Exception as e:
            raise Exception(f"处理PDF文档时出错: {str(e)}")

    def _process_text(self):
        """处理文本文档"""
        try:
            section = self._create_default_section()
            text_content = self.file_content.decode('utf-8')
            
            content = DocumentContent(
                self.document_id,
                section['_id'],
                ContentType.TEXT
            )
            content.set_text_content(text_content)
            content.data['_id'] = str(uuid.uuid4())
            content.data['order'] = 10
            document_contents.insert_one(content.data)
        except Exception as e:
            raise Exception(f"处理文本文档时出错: {str(e)}") 